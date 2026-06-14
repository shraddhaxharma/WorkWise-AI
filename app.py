import os
import streamlit as st
import shutil

from langchain.schema import Document
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import (
    GoogleGenerativeAIEmbeddings,
    ChatGoogleGenerativeAI,
)
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate

from pypdf import PdfReader
from docx import Document as DocxDocument

# Load API Key
load_dotenv()

if not os.getenv("GOOGLE_API_KEY"):
    st.error("GOOGLE_API_KEY not found!")
    st.stop()

st.set_page_config(
    page_title="WorkWise AI",
    page_icon="📚",
    layout="wide",
)

st.title("WorkWise AI")

st.caption(
    "Your workplace knowledge assistant — instantly find answers from SOPs, knowledge base articles, policies and process documents."
)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Read Documents

def get_documents(docs):

    documents = []

    for file in docs:

        filename = file.name

        # PDF

        if filename.lower().endswith(".pdf"):

            pdf_reader = PdfReader(file)

            for page_no, page in enumerate(pdf_reader.pages):

                page_text = page.extract_text()

                if page_text:

                    documents.append(
                        Document(
                            page_content=page_text,
                            metadata={
                                "source": filename,
                                "page": page_no + 1,
                            },
                        )
                    )

        # TXT

        elif filename.lower().endswith(".txt"):

            text = file.read().decode("utf-8")

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": filename,
                            "page": 1,
                        },
                    )
                )

        # DOCX

        elif filename.lower().endswith(".docx"):

            doc = DocxDocument(file)

            text = "\n".join(
                para.text
                for para in doc.paragraphs
                if para.text.strip()
            )

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": filename,
                            "page": 1,
                        },
                    )
                )

    return documents

# Split Text

def get_text_chunks(documents):

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
    )

    return splitter.split_documents(documents)

# Create Vector Store

def get_vector_store(text_chunks):

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001"
    )

    if os.path.exists("faiss_index"):
        shutil.rmtree("faiss_index")

    vector_store = FAISS.from_documents(
        text_chunks,
        embedding=embeddings,
    )

    vector_store.save_local("faiss_index")

# Answer Question

def user_input(question):

    if not os.path.exists("faiss_index/index.faiss"):

        st.warning(
            "⚠️ Please process your documents before asking a question."
        )

        return

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001"
    )

    db = FAISS.load_local(
        "faiss_index",
        embeddings,
        allow_dangerous_deserialization=True,
    )

    docs = db.similarity_search(
        question,
        k=3,
    )

    model = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0.3,
        google_api_key=os.getenv("GOOGLE_API_KEY"),
    )

    prompt_template = """
You are an AI Knowledge Assistant.

Use ONLY the context provided below.

If the answer is not found in the context, respond exactly with:

❌ I could not find this information in the uploaded documents.

Do not use your own knowledge.
Do not guess.
Do not make up information.

Context:
{context}

Question:
{question}

Answer:
"""

    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["context", "question"],
    )

    chain = load_qa_chain(
        model,
        chain_type="stuff",
        prompt=prompt,
    )

    with st.spinner("🤖 Generating answer..."):

        response = chain.run(
            input_documents=docs,
            question=question,
        )

    st.session_state.chat_history.append(
        {
            "role": "user",
            "content": question,
        }
    )

    st.session_state.chat_history.append(
        {
            "role": "assistant",
            "content": response,
        }
    )

    with st.chat_message("assistant"):
        st.markdown(response)

    with st.expander("📄 Sources Used"):

        for i, doc in enumerate(docs):

            st.markdown(f"### 📑 Source {i+1}")

            st.markdown(
                f"**📄 File:** {doc.metadata.get('source','Unknown')}"
            )

            st.markdown(
                f"**📃 Page:** {doc.metadata.get('page','Unknown')}"
            )

            st.write(doc.page_content)

            st.divider()

# Sidebar

with st.sidebar:

    st.header("📂 Repository")

    st.info(
        """
Upload workplace documents & WorkWise AI will retrieve relevant information and provide accurate, source-backed answers in seconds.
"""
    )

    docs = st.file_uploader(
        "Upload Documents",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
    )

    if st.button("Process Documents"):

        if not docs:

            st.warning(
                "⚠️ Please upload at least one document."
            )

        else:

            with st.spinner("Processing documents..."):

                documents = get_documents(docs)

                if not documents:

                    st.warning(
                        "⚠️ No readable text found in the uploaded documents."
                    )

                    st.stop()

                text_chunks = get_text_chunks(documents)

                get_vector_store(text_chunks)

            st.success(
                """
✅ Processed successfully.

You can now search your uploaded documents.
"""
            )

    st.divider()

    if st.button("🗑️ Clear Conversation"):

        st.session_state.chat_history = []

        if os.path.exists("faiss_index"):

            shutil.rmtree("faiss_index")

        st.success(
            "Conversation and knowledge base cleared."
        )

        st.rerun()

# Chat History

for message in st.session_state.chat_history:

    with st.chat_message(message["role"]):

        st.markdown(message["content"])

# Question Box

question = st.chat_input(
    "Ask a question about your documents..."
)

if question:

    with st.chat_message("user"):

        st.markdown(question)

    user_input(question)